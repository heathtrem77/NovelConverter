import sys

from PyQt5.QtGui import *
from PyQt5.QtCore import *
from PyQt5.QtWidgets import *

from docx import Document

from NovelConverter_layout import *

exception = ['【','［','「','『','〈','〝']
restoreEpn = ['　']
isModeFormat = True
txtType = 'Notepad (*.txt)'
wordType = 'Microsoft Word (*.doc *docx)'
initFilter = txtType

class NCWindow(QDialog, Ui_NovelConverter):
    def __init__(self):
        QDialog.__init__(self)
        self.setupUi(self)
        self.setWindowFlags(self.windowFlags() | QtCore.Qt.WindowMinimizeButtonHint | QtCore.Qt.WindowMaximizeButtonHint | QtCore.Qt.WindowSystemMenuHint)
        self.show()

        self.inputPlace.textChanged.connect(self.convertOriginalTextToUploadFormat)
        self.pushButton.clicked.connect(self.modeChanger)
        self.importButton.clicked.connect(self.fileImport)
        self.exportButton.clicked.connect(self.fileExport)
    
    def convertOriginalTextToUploadFormat(self):
        oldText = self.inputPlace.toPlainText().splitlines()
        length = len(oldText)
        new = ''
        lineChange = '\n\n'

        for i, line in enumerate(oldText):
            if i == length - 1:
                lineChange = ''
            
            if line:
                if line[0] in exception:
                    new += line + lineChange
                else:
                    new += '　' + line + lineChange
            else:
                new += lineChange
        
        self.outputPlace.clear()
        self.outputPlace.setPlainText(new)

    def convertUploadFormatToOriginalText(self):
        oldText = self.inputPlace.toPlainText().splitlines()
        length = len(oldText)
        new = ''
        lineChange = '\n'
        cnt = 0

        for i, line in enumerate(oldText):
            if i == length - 1:
                lineChange = ''
            
            if line:
                cnt = 0
                if line[0] in restoreEpn:
                    new += line[1:] + lineChange
                else:
                    new += line + lineChange
            else:
                cnt += 1
                if cnt > 1:
                    new += lineChange
                    cnt = 0
        
        self.outputPlace.clear()
        self.outputPlace.setPlainText(new)
    
    def modeChanger(self):
        global isModeFormat
        _translate = QtCore.QCoreApplication.translate
        if isModeFormat:
            self.inputPlace.textChanged.connect(self.convertUploadFormatToOriginalText)
            self.pushButton.setText(_translate("NovelConverter", "Go to format mode"))
            isModeFormat = False
            self.outputPlace.clear()
        else:
            self.inputPlace.textChanged.connect(self.convertOriginalTextToUploadFormat)
            self.pushButton.setText(_translate("NovelConverter", "Go to restore mode"))
            isModeFormat = True
            self.outputPlace.clear()
    
    def fileImport(self):
        # support *.txt and *.docx
        global initFilter
        getFile = QFileDialog.getOpenFileName(self, '', '', f'{txtType};; {wordType}', initFilter)

        if getFile[0]:
            if getFile[1] == txtType:
                with open(getFile[0], 'r', encoding='utf-8') as f:
                    self.inputPlace.setPlainText(f.read())
            else:
                f = Document(getFile[0]).paragraphs
                length = len(f)
                new = ''
                lineChange = '\n'

                for i in range(length):
                    new += f[i].text
                    if i < length-1:
                        new += lineChange
                self.inputPlace.setPlainText(new)
            initFilter = getFile[1]
    
    def fileExport(self):
        # support *.txt and *.docx
        global initFilter
        saveFile = QFileDialog.getSaveFileName(self, 'Save File', '', f'{txtType};; {wordType}', initFilter)
        targetText = self.outputPlace.toPlainText()

        if saveFile[0]:
            if saveFile[1] == txtType:
                with open(saveFile[0], 'w', encoding='utf-8') as f:
                    f.write(targetText)
            else:
                f = Document()
                targetText = targetText.splitlines()

                for i, line in enumerate(targetText):
                    f.add_paragraph(line)
                f.save(saveFile[0])
            initFilter = saveFile[1]

if __name__ == '__main__':
    app = QApplication(sys.argv)
    mywindow = NCWindow()
    mywindow.show()
    app.exec_()